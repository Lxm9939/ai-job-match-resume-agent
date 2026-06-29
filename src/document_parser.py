"""Resume file parsing for txt, pdf and docx uploads."""

from __future__ import annotations

from io import BytesIO
from typing import BinaryIO

from docx import Document
from pypdf import PdfReader

from src.utils.text_utils import normalize_text


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def parse_document(file_name: str, content: bytes) -> str:
    """Parse uploaded resume bytes into plain text."""

    suffix = "." + file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")
    if suffix == ".txt":
        return _parse_txt(content)
    if suffix == ".pdf":
        return _parse_pdf(content)
    if suffix == ".docx":
        return _parse_docx(content)
    return ""


def parse_uploaded_file(uploaded_file: BinaryIO) -> str:
    """Parse a Streamlit UploadedFile-like object."""

    file_name = getattr(uploaded_file, "name", "resume")
    content = uploaded_file.read()
    return parse_document(file_name, content)


def _parse_txt(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
        try:
            return normalize_text(content.decode(encoding))
        except UnicodeDecodeError:
            continue
    return normalize_text(content.decode("utf-8", errors="ignore"))


def _parse_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return normalize_text("\n".join(pages))


def _parse_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return normalize_text("\n".join(parts))

