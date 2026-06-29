"""Export workflow results to shareable report files."""

from __future__ import annotations

from io import BytesIO
from typing import Iterable, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from src.schemas.models import WorkflowResult


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def export_workflow_result_to_docx(result: WorkflowResult) -> bytes:
    """Build a Word report from a completed workflow result."""

    document = Document()
    _configure_document(document)

    title = document.add_heading("AI 秋招岗位匹配与简历优化报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_job_summary(document, result)
    _add_jd_analysis(document, result)
    _add_evidence_table(document, result)
    _add_keyword_coverage(document, result)
    _add_score_breakdown(document, result)
    _add_resume_optimization(document, result)
    _add_outreach_messages(document, result)
    _add_final_conclusion(document, result)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)


def _add_job_summary(document: Document, result: WorkflowResult) -> None:
    jd = result.jd
    score = result.score
    document.add_heading("一、岗位摘要", level=1)
    _add_key_value_table(
        document,
        [
            ("岗位名称", jd.job_title or "未识别"),
            ("公司", jd.company or "未识别"),
            ("地点", jd.location or "未识别"),
            ("总匹配分", f"{score.total_score}/100"),
            ("核心结论", score.summary or "暂无结论"),
        ],
    )


def _add_jd_analysis(document: Document, result: WorkflowResult) -> None:
    jd = result.jd
    document.add_heading("二、JD 解析结果", level=1)
    _add_bullets(document, "岗位职责", jd.responsibilities)
    _add_bullets(document, "硬技能", jd.hard_skills)
    _add_bullets(document, "软技能", jd.soft_skills)
    _add_bullets(document, "工具栈", jd.tools)
    _add_bullets(document, "业务关键词", jd.business_keywords)
    _add_bullets(document, "隐含能力", jd.implicit_capabilities)
    _add_key_value_table(
        document,
        [
            ("学历要求", jd.education_requirement or "未识别"),
            ("经验要求", jd.experience_requirement or "未识别"),
        ],
    )


def _add_evidence_table(document: Document, result: WorkflowResult) -> None:
    document.add_heading("三、简历证据匹配表", level=1)
    rows = [
        [item.requirement, item.resume_evidence, item.strength, item.suggested_expression]
        for item in result.evidence_matches
    ]
    _add_table(document, ["岗位要求", "简历证据", "证据强度", "建议补充表达"], rows)


def _add_keyword_coverage(document: Document, result: WorkflowResult) -> None:
    coverage = result.keyword_coverage
    document.add_heading("四、关键词覆盖分析", level=1)
    document.add_paragraph(f"关键词强覆盖率：{coverage.coverage_rate:.0%}")
    if coverage.notes:
        document.add_paragraph(coverage.notes)
    _add_bullets(document, "已覆盖关键词", coverage.covered_keywords)
    _add_bullets(document, "弱覆盖关键词", coverage.weak_keywords)
    _add_bullets(document, "未覆盖关键词", coverage.missing_keywords)


def _add_score_breakdown(document: Document, result: WorkflowResult) -> None:
    document.add_heading("五、匹配评分", level=1)
    rows = [[item.name, f"{item.weight:.0%}", f"{item.score}", item.reason] for item in result.score.categories]
    _add_table(document, ["评分维度", "权重", "得分", "原因"], rows)
    _add_bullets(document, "优势", result.score.strengths)
    _add_bullets(document, "风险与短板", result.score.risks)


def _add_resume_optimization(document: Document, result: WorkflowResult) -> None:
    document.add_heading("六、简历优化建议", level=1)
    document.add_paragraph("原则：不编造经历，不虚构项目，不添加用户没做过的内容。")
    rows = [
        [item.original_bullet, item.optimized_bullet, item.rationale, item.risk_note]
        for item in result.optimization_suggestions
    ]
    _add_table(document, ["修改前", "修改后", "理由", "风险提示"], rows)


def _add_outreach_messages(document: Document, result: WorkflowResult) -> None:
    outreach = result.outreach
    document.add_heading("七、投递话术", level=1)
    _add_key_value_table(
        document,
        [
            ("Boss 直聘打招呼", outreach.boss_zhipin),
            ("邮件投递正文", outreach.email_body),
            ("LinkedIn 私信", outreach.linkedin_dm),
            ("内推请求话术", outreach.referral_request),
            ("面试自我介绍初稿", outreach.interview_intro),
        ],
    )


def _add_final_conclusion(document: Document, result: WorkflowResult) -> None:
    document.add_heading("八、最终结论", level=1)
    document.add_paragraph(result.score.summary or "暂无最终结论。")
    if result.score.risks:
        document.add_paragraph("建议优先补强：")
        for risk in result.score.risks[:5]:
            document.add_paragraph(risk, style="List Bullet")


def _add_bullets(document: Document, title: str, values: Sequence[str]) -> None:
    document.add_paragraph(title, style="Intense Quote")
    if not values:
        document.add_paragraph("未识别")
        return
    for value in values:
        document.add_paragraph(_safe_text(value), style="List Bullet")


def _add_key_value_table(document: Document, rows: Sequence[tuple[str, str]]) -> None:
    _add_table(document, ["项目", "内容"], rows)


def _add_table(document: Document, headers: Sequence[str], rows: Iterable[Sequence[object]]) -> None:
    rows = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    if not rows:
        cells = table.add_row().cells
        cells[0].text = "暂无"
        for index in range(1, len(headers)):
            cells[index].text = ""
    else:
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                if index < len(cells):
                    cells[index].text = _safe_text(value)
    document.add_paragraph()


def _safe_text(value: object) -> str:
    text = "" if value is None else str(value)
    return text.replace("\r\n", "\n").replace("\r", "\n")
