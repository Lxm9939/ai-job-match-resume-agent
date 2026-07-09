from __future__ import annotations

import re
import zipfile
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = (
    PROJECT_ROOT
    / "docs"
    / "deliverables"
    / "AI岗位搜索匹配评分与求职优化助手_项目计划书.docx"
)
SCREENSHOT_DIR = PROJECT_ROOT / "docs" / "screenshots"
WORKFLOW_MD = PROJECT_ROOT / "docs" / "workflow.md"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 98, 110)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D7DBE2"


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "一、项目概述", 1)
    add_para(
        doc,
        "本项目是一个面向不同求职阶段候选人的网页端求职辅助工具。用户可以上传简历文件或粘贴简历文本，"
        "再输入岗位 JD 和目标岗位方向，系统通过 8 个模块化 Agent 完成岗位解析、简历解析、证据匹配、"
        "关键词覆盖、匹配评分、简历优化、投递话术生成和最终报告汇总。"
    )
    add_para(
        doc,
        "当前版本已经实现可运行的 Streamlit MVP，支持 mock 模式本地演示，也支持在配置 OpenAI API Key 后进行 LLM 增强解析。"
        "输出结果通过 7 个 tab 展示，并支持 Markdown 和 Word .docx 两种报告导出方式。"
    )
    add_key_value_table(
        doc,
        [
            ("项目名称", "AI 岗位搜索、匹配评分与求职优化助手"),
            ("项目目录", str(PROJECT_ROOT)),
            ("应用形态", "Python + Streamlit 网页应用"),
            ("核心架构", "模块化 Agent 工作流 + Pydantic 数据模型 + 可选 OpenAI API"),
            ("当前状态", "已完成 MVP、文档、测试用例、截图和报告导出功能"),
        ],
    )

    add_heading(doc, "二、背景痛点", 1)
    add_bullets(
        doc,
        [
            "不同岗位的 JD 信息复杂，岗位职责、任职要求、工具栈、业务关键词和隐含能力常常混在一起。",
            "候选人难以快速判断简历经历与具体岗位要求之间的证据关系。",
            "简历关键词覆盖不足时，容易在招聘方快速浏览或 ATS 初筛环节失分。",
            "批量投递时，手动改简历和编写多渠道投递话术效率较低。",
            "直接让 LLM 改写简历存在编造经历风险，需要明确约束和可解释流程。",
        ],
    )

    add_heading(doc, "三、目标用户", 1)
    add_bullets(
        doc,
        [
            "应届生、实习求职者、海外留学生和回国求职用户。",
            "社招候选人、在职跳槽用户和转行到数据、AI、产品、分析岗位的人。",
            "申请 AI 应用、Data Agent、数据分析、商业分析、BI 分析和产品经理等岗位的人。",
            "需要批量比较岗位、优化简历并准备面试的求职者。",
        ],
    )

    add_heading(doc, "四、核心功能", 1)
    feature_rows = [
        ("简历输入", "支持上传 .txt / .pdf / .docx 简历，也支持直接粘贴简历文本。"),
        ("JD 输入", "支持粘贴岗位 JD，并填写目标岗位方向。"),
        ("JD 解析", "提取岗位名称、公司、地点、职责、硬技能、软技能、工具栈、业务关键词、学历经验要求和隐含能力。"),
        ("简历解析", "提取教育背景、项目经历、实习经历、技能栈、成果指标、关键词和可迁移能力。"),
        ("证据匹配", "输出“岗位要求 - 简历证据 - 证据强度 - 建议补充表达”。"),
        ("关键词覆盖", "输出已覆盖、弱覆盖、未覆盖关键词和覆盖率说明。"),
        ("匹配评分", "按技能 30%、项目经历 25%、关键词覆盖 20%、岗位职责 15%、教育/背景 10% 计算总分。"),
        ("简历优化", "基于已有 bullet 给出修改建议，明确不编造经历、不虚构项目。"),
        ("投递话术", "生成 Boss 直聘、邮件、LinkedIn、内推请求和面试自我介绍初稿。"),
        ("报告导出", "支持 Markdown 报告下载和 Word .docx 报告导出。"),
    ]
    add_table(doc, ["功能模块", "当前实现说明"], feature_rows, [1800, 7560])

    add_heading(doc, "五、Agent 工作流", 1)
    add_para(
        doc,
        "项目实际代码中包含 8 个 Agent 文件，并由 src/workflow.py 编排为端到端分析流程。"
        "docs/workflow.md 中的 Mermaid 流程图如下；由于 Word 默认不渲染 Mermaid，这里保留代码块，方便复制到 GitHub 或 Mermaid 编辑器查看。"
    )
    add_code_block(doc, extract_mermaid(WORKFLOW_MD))
    agent_rows = [
        ("JD 解析 Agent", "src/agents/jd_parser_agent.py", "JD 文本、目标岗位方向", "JDAnalysis"),
        ("简历解析 Agent", "src/agents/resume_parser_agent.py", "简历纯文本", "ResumeAnalysis"),
        ("简历证据提取 Agent", "src/agents/evidence_agent.py", "JDAnalysis、ResumeAnalysis", "EvidenceMatch 列表"),
        ("关键词覆盖 Agent", "src/agents/keyword_agent.py", "JDAnalysis、ResumeAnalysis", "KeywordCoverage"),
        ("匹配评分 Agent", "src/agents/scoring_agent.py", "JD、简历、证据、关键词覆盖", "ScoreBreakdown"),
        ("简历优化 Agent", "src/agents/resume_optimizer_agent.py", "JD、简历", "OptimizationSuggestion 列表"),
        ("话术生成 Agent", "src/agents/outreach_agent.py", "JD、简历、评分、目标岗位方向", "OutreachMessages"),
        ("最终报告 Agent", "src/agents/report_agent.py", "全部中间结果", "FinalReport"),
    ]
    add_table(doc, ["Agent", "代码文件", "输入", "输出"], agent_rows, [1900, 2850, 2500, 2110])

    add_heading(doc, "六、技术架构", 1)
    add_table(
        doc,
        ["层级", "已实现技术/模块", "说明"],
        [
            ("前端交互", "Streamlit / app.py", "提供简历上传、文本输入、JD 输入、分析按钮和 7 个结果 tab。"),
            ("工作流编排", "src/workflow.py", "串联 8 个 Agent，输出 WorkflowResult。"),
            ("结构化模型", "src/schemas/models.py", "使用 Pydantic 定义 JD、简历、评分、话术和报告等模型。"),
            ("文档解析", "src/document_parser.py", "支持 .txt、.pdf、.docx 简历解析。"),
            ("LLM 封装", "src/llm_client.py", "封装 OpenAI API，并提供 mock 回退。"),
            ("规则工具", "src/utils/text_utils.py", "提供关键词抽取、文本清洗、表格生成等工具函数。"),
            ("报告导出", "src/report_exporter.py", "使用 python-docx 生成 Word 报告。"),
            ("测试", "tests/test_workflow.py", "覆盖 mock 工作流和 Word 导出。"),
        ],
        [1500, 2700, 5160],
    )
    add_callout(
        doc,
        "当前版本不依赖真实 API Key 即可演示。LLM_MODE=mock 时，系统使用本地规则和关键词库完成分析；配置 OpenAI API Key 后可切换为增强模式。"
    )

    add_heading(doc, "七、数据输入输出", 1)
    add_table(
        doc,
        ["类别", "内容", "当前支持情况"],
        [
            ("输入", "简历文件", ".txt / .pdf / .docx"),
            ("输入", "简历文本", "直接粘贴到 Streamlit 文本框"),
            ("输入", "岗位 JD", "直接粘贴岗位职责、要求、加分项等"),
            ("输入", "目标岗位方向", "例如 AI 产品经理、数据分析师、商业分析师、数据产品经理、AI Agent 产品经理"),
            ("输出", "页面结果", "7 个 tab：JD 解析、证据表、关键词、评分、优化、话术、最终报告"),
            ("输出", "报告文件", "Markdown .md 和 Word .docx"),
            ("输出", "测试/示例", "examples/sample_resume.txt、examples/sample_jd.txt、docs/test_cases.md"),
        ],
        [1200, 2800, 5360],
    )

    add_heading(doc, "八、测试方案", 1)
    add_para(
        doc,
        "当前项目已包含 pytest 自动测试和 docs/test_cases.md 手动测试矩阵。自动测试覆盖 mock 模式端到端工作流，以及 Word .docx 报告导出是否生成有效文件。"
    )
    add_code_block(doc, "pytest")
    add_table(
        doc,
        ["测试类别", "已覆盖内容", "检查点"],
        [
            ("自动测试", "mock 模式工作流", "JD 解析、简历解析、证据匹配、评分、最终报告均有结果"),
            ("自动测试", "Word 报告导出", ".docx 字节以 ZIP/OOXML 格式开头，文件大小有效"),
            ("手动测试", "文本输入和文件上传", "简历粘贴、.txt/.pdf/.docx 上传均可触发分析"),
            ("手动测试", "缺失输入", "简历为空或 JD 为空时显示清晰错误提示"),
            ("手动测试", "Mock/OpenAI 模式", "无 Key 可演示，有 Key 可增强，失败可回退"),
            ("UI 启动检查", "Streamlit 本地启动", "页面可访问并返回 HTTP 200"),
        ],
        [1600, 3500, 4260],
    )

    add_heading(doc, "九、风险与解决方案", 1)
    add_table(
        doc,
        ["风险", "影响", "当前解决方案"],
        [
            ("LLM 输出不稳定", "结构化解析失败", "LLMClient 统一封装，失败回退到本地规则结果。"),
            ("无 API Key", "无法在线调用模型", "默认 mock 模式可完整本地演示。"),
            ("简历优化编造内容", "求职真实性风险", "Agent 提示和 UI 明确要求不编造经历、不虚构项目。"),
            ("PDF/DOCX 解析质量波动", "简历文本可能不完整", "提供文本粘贴作为兜底输入。"),
            ("评分被误解为绝对结论", "用户过度依赖单一分数", "输出分项原因、优势和风险解释。"),
            ("截图素材维护成本", "README 展示可能过期", "docs/screenshots/README.md 维护截图文件清单。"),
        ],
        [2200, 2300, 4860],
    )

    add_heading(doc, "十、项目亮点", 1)
    add_bullets(
        doc,
        [
            "真实可运行的 Streamlit MVP，而不是纯文档或伪原型。",
            "将求职分析拆成 8 个可解释 Agent，便于后续扩展和调试。",
            "Mock 模式保证无 API Key 也能演示，适合 GitHub 和作品集展示。",
            "支持 Markdown 和 Word 报告导出，形成可保存、可复盘的交付物。",
            "保留 Mermaid 工作流、测试用例、截图清单和 README 展示材料，项目完整度较高。",
            "强调“不编造经历”的简历优化原则，贴近真实求职场景。",
        ],
    )

    add_heading(doc, "十一、GitHub 展示说明", 1)
    add_para(
        doc,
        "README.md 已改为作品集展示版本，包含项目简介、背景痛点、核心功能、Agent 工作流图、页面截图区域、项目结构、快速开始、环境变量、Mock 模式、测试说明和后续迭代计划。"
    )
    screenshot_rows = []
    screenshot_map = [
        ("首页和输入区", "home.png"),
        ("JD 解析 tab", "jd_parser.png"),
        ("简历证据匹配 tab", "evidence.png"),
        ("关键词覆盖 tab", "keywords.png"),
        ("匹配评分 tab", "score.png"),
        ("简历优化 tab", "resume_optimization.png"),
        ("投递话术 tab", "outreach.png"),
        ("最终报告 tab", "final_report.png"),
    ]
    for label, filename in screenshot_map:
        path = SCREENSHOT_DIR / filename
        status = "已存在" if path.exists() else "缺失"
        screenshot_rows.append((label, f"docs/screenshots/{filename}", status))
    add_table(doc, ["截图内容", "文件路径", "当前状态"], screenshot_rows, [2400, 5000, 1960])

    add_heading(doc, "关键页面截图", 2)
    add_screenshot(doc, "首页和输入区", SCREENSHOT_DIR / "home.png")
    add_screenshot(doc, "匹配评分", SCREENSHOT_DIR / "score.png")
    add_screenshot(doc, "简历优化建议", SCREENSHOT_DIR / "resume_optimization.png")

    add_heading(doc, "十二、后续迭代计划", 1)
    add_para(doc, "以下内容为后续计划，当前版本尚未实现，不作为 MVP 已交付能力描述。")
    add_bullets(
        doc,
        [
            "PDF 报告导出：当前已实现 Word .docx，PDF 导出后续可通过 LibreOffice 或专用 PDF 管线补充。",
            "批量 JD 匹配排序：支持同时输入多个岗位并按匹配度排序。",
            "岗位关键词库：为 AI 产品、数据分析、商业分析、数据产品等方向维护专属关键词库。",
            "多版本简历对比：记录不同版本简历的匹配分和优化建议。",
            "更严格 JSON Schema 校验：增强 LLM 输出修复、重试和格式约束。",
            "UI 自动化测试：增加 Playwright 或 Streamlit 测试脚本，覆盖关键页面和下载链路。",
            "在线部署：后续可部署到 Streamlit Community Cloud、Hugging Face Spaces 或其他平台。",
        ],
    )

    add_heading(doc, "十三、简历项目描述", 1)
    add_callout(
        doc,
        "AI 岗位搜索、匹配评分与求职优化助手：基于 Python、Streamlit 和多 Agent 工作流构建通用求职辅助应用，"
        "支持简历上传/粘贴、JD 解析、简历证据匹配、关键词覆盖分析、可解释匹配评分、简历 bullet 优化、"
        "多渠道投递话术生成和 Markdown/Word 报告导出；封装 OpenAI API 并提供无 Key 可运行的 mock 模式，"
        "使用 Pydantic 管理结构化输出，pytest 覆盖端到端工作流和报告导出。"
    )

    add_footer(doc)
    doc.save(OUTPUT_PATH)
    assert_valid_docx(OUTPUT_PATH)
    print(str(OUTPUT_PATH).encode("unicode_escape").decode("ascii"))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing = 1.0


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "AI 岗位搜索、匹配评分与求职优化助手 | 项目计划书"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("作品集项目计划书")
    kr.font.name = "Microsoft YaHei"
    kr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    kr.font.size = Pt(12)
    kr.font.bold = True
    kr.font.color.rgb = BLUE

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("AI 岗位搜索、匹配评分与求职优化助手")
    tr.font.name = "Microsoft YaHei"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    tr.font.size = Pt(26)
    tr.font.bold = True
    tr.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("基于 Streamlit 与多 Agent 工作流的求职辅助工具")
    sr.font.name = "Microsoft YaHei"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    sr.font.size = Pt(14)
    sr.font.color.rgb = MUTED

    add_key_value_table(
        doc,
        [
            ("项目状态", "已实现可运行 MVP"),
            ("核心能力", "简历上传/粘贴、JD 输入、8 个 Agent、7 个结果 tab、Mock 模式、Markdown/Word 报告导出"),
            ("技术栈", "Python、Streamlit、Pydantic、python-docx、OpenAI API 封装"),
            ("生成日期", date.today().isoformat()),
        ],
    )
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    sections = [
        "一、项目概述",
        "二、背景痛点",
        "三、目标用户",
        "四、核心功能",
        "五、Agent 工作流",
        "六、技术架构",
        "七、数据输入输出",
        "八、测试方案",
        "九、风险与解决方案",
        "十、项目亮点",
        "十一、GitHub 展示说明",
        "十二、后续迭代计划",
        "十三、简历项目描述",
    ]
    for item in sections:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_fill(cell, CALLOUT_FILL)
    set_cell_margins(cell, 140, 140, 180, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_fill(cell, "F7F7F7")
    set_cell_margins(cell, 100, 100, 140, 140)
    for idx, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.style = "Code Block"
        p.add_run(line)
    doc.add_paragraph()


def add_key_value_table(doc: Document, rows: Sequence[tuple[str, str]]) -> None:
    add_table(doc, ["项目", "内容"], rows, [1800, 7560])


def add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[object]], widths: Sequence[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0]
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = str(header)
        set_cell_fill(cell, LIGHT_FILL)
        set_cell_margins(cell, 80, 80, 120, 120)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = NAVY
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx], 80, 80, 120, 120)
    doc.add_paragraph()


def add_screenshot(doc: Document, caption: str, path: Path) -> None:
    if not path.exists():
        add_para(doc, f"{caption}：截图文件缺失，预期路径 {path.relative_to(PROJECT_ROOT)}。")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(str(path), width=Inches(6.2))


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "AI 岗位搜索、匹配评分与求职优化助手项目计划书"
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = MUTED


def extract_mermaid(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    match = re.search(r"```mermaid\s*(.*?)```", text, re.S)
    if not match:
        return "flowchart TD\n    A[未找到 Mermaid 工作流]"
    return match.group(1).strip()


def set_table_width(table, widths: Sequence[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(widths[idx]))
                tc_w.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int, bottom: int, start: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def assert_valid_docx(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
    required = {"[Content_Types].xml", "word/document.xml"}
    missing = required - names
    if missing:
        raise RuntimeError(f"Invalid docx, missing: {missing}")


if __name__ == "__main__":
    main()
